import re
from collections import defaultdict

from openpyxl import load_workbook
import docx
from tibetan_sort.tibetan_sort import TibetanSort

bo_sort = TibetanSort()


def get_sentences_after(props, length):
    if length == len(props):
        return props[length-1]
    else:
        subprops = get_sentences_after(props, length + 1)
        res = []
        for thisprop in props[length-1]:
            for thissubprop in subprops:
                res.append(thisprop + ' ' + thissubprop)
        return res


def parse_sheets(xlsx):
    wb = load_workbook(xlsx)
    sheets = dict()
    for sheet in wb.worksheets:
        sheet_name = sheet.title
        rows = []
        meta = []
        for i, row in enumerate(sheet):
            values = [r.value for r in row]
            if i == 0:
                meta = values
            elif [v for v in values if v]:
                rows.append(values)
            else:
                break
        sheets[sheet_name] = (meta, rows)
    return sheets


def get_chunk_idx(row):
    chunk_idx = []
    start = 0
    for i, r in enumerate(row):
        if not r:
            end = i - 1
            chunk_idx.append((start, end))
            start = i + 1
    chunk_idx.append((start, i))
    return chunk_idx


def process_chunks(sheets):
    total_chunks = dict()
    for name, v in sheets.items():
        meta, rows = v
        idx = get_chunk_idx(rows[0])

        chunks = []
        for a, b in idx:
            # generate all versions
            chunk_versions = [r[a:b+1] for r in rows]
            chunk_versions = [''.join([c for c in chunk if c]) for chunk in chunk_versions]
            # remove duplicates
            chunk_versions = list(set([c for c in chunk_versions if c]))

            # add optionality
            required = True if [m for m in meta[a:b+1] if m] else False
            if not required:
                chunk_versions.append('')

            chunks.append(chunk_versions)
        total_chunks[name] = chunks
    return total_chunks


def export_docx(sentences, out_file):
    doc = docx.Document()
    for section, i in sentences.items():
        sents, orig = i
        doc.add_heading(f'{section} {orig}', level=1)
        doc.add_paragraph('')
        for size in sorted(sents.keys()):
            doc.add_heading(f'དུམ་བུ་ {size}ཡོད་པ།', level=2)
            for s in sents[size]:
                doc.add_paragraph(s, style='List Bullet 2')
    doc.save(out_file)


def sort_sents(sents):
    orig = sents[0]
    # group by size
    by_size = defaultdict(list)
    for sent in sents:
        chunks = [s for s in sent.split(' ') if s]
        l = len(chunks)
        by_size[l].append(sent)
    # sort groups
    for k, v in by_size.items():
        by_size[k] = bo_sort.sort_list(v)
    return by_size, orig


def gen_alt_sentences(in_file, out_file):
    """

    :param in_file: xlsx file, one sentence per sheet.
                    chunks are separated by an empty column.
                    1st row has info about which chunks are required: 1 or more cells of chunk is not empty if required
    :param out_file: docx file
    """
    sheets = parse_sheets(in_file)
    chunks = process_chunks(sheets)

    sentences = dict()
    for name, parts in chunks.items():
        sents = get_sentences_after(parts, 1)
        sents = [re.sub(r'\s+', ' ', s) for s in sents]  # single spaces
        sents = sort_sents(sents)
        sentences[name] = sents

    export_docx(sentences, out_file)


if __name__ == '__main__':
    in_file = 'input/ཁ་འདོན།.xlsx'
    out_file = 'output/ཁ་འདོན།_sents.docx'
    gen_alt_sentences(in_file, out_file)
